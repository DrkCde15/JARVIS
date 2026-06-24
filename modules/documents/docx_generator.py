import os
from pathlib import Path
from typing import Optional

from modules.documents.base import ensure_output_dir, safe_filename


def generate_docx(
    title: str,
    content: list[dict],
    filename: Optional[str] = None,
    author: str = "JARVIS",
    template_path: Optional[str] = None,
) -> str:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    doc.add_heading(title, level=0)

    for block in content:
        block_type = block.get("type", "paragraph")
        text = block.get("text", "")

        if block_type == "heading1":
            doc.add_heading(text, level=1)
        elif block_type == "heading2":
            doc.add_heading(text, level=2)
        elif block_type == "heading3":
            doc.add_heading(text, level=3)
        elif block_type == "paragraph":
            p = doc.add_paragraph(text)
            if block.get("style") == "bullet":
                p.style = doc.styles["List Bullet"]
            elif block.get("style") == "number":
                p.style = doc.styles["List Number"]
        elif block_type == "code":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 30, 30)
        elif block_type == "table":
            table_data = block.get("data", [])
            if table_data:
                rows_count = len(table_data)
                cols_count = max(len(r) for r in table_data) if table_data else 0
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        if j < cols_count:
                            table.cell(i, j).text = str(cell_text)
        elif block_type == "page_break":
            doc.add_page_break()

    doc.core_properties.author = author

    output_dir = ensure_output_dir("docx")
    name = filename or safe_filename(title)
    if not name.endswith(".docx"):
        name += ".docx"

    output_path = output_dir / name
    doc.save(str(output_path))
    return str(output_path)
